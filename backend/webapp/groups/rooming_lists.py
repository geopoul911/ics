
from accounts.models import UserProfile, User
from webapp.models import (
    GroupTransfer,
    History,
    Contact,
    Hotel,
    DocRoomingList,
    Agent,
    TravelDay
)
from django.views.decorators.csrf import csrf_exempt
from rest_framework.authtoken.models import Token
from rest_framework import generics
from rest_framework.response import Response
import datetime
from django.core.mail import EmailMultiAlternatives
from webapp.serializers import (
    RoomingListSerializer,
    GroupSerializer,
)
from accounts.permissions import (
    can_update,
)

from bs4 import BeautifulSoup
from docx import Document

from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import RGBColor
import os
from pathlib import Path

"""
    # Rooming lists

    - ChangeRoomingListRoomText
    - ChangeRoomingListRecipient
    - ChangeRoomingListRoomDescr
    - ChangeRoomingListMealDescr
    - ChangeRoomingListNote
    - EditMassRoomingListFields
    - SendAllRoomingLists
    - SendRoomingList

"""


def days_are_consecutive(days):
    # Sort the days to ensure they are in chronological order
    sorted_days = sorted(days)

    # Check if each day is exactly one day after the previous day
    for i in range(1, len(sorted_days)):
        if sorted_days[i] - sorted_days[i-1] != datetime.timedelta(days=1):
            return False
    return True


BASE_DIR = Path(__file__).resolve().parent.parent.parent


# Get user from token
def get_user(token):
    return Token.objects.get(key=token).user


def update_roominglists(group_transfer):

    """Updates the rooming lists.
    Usually called when hotels change.
    """
    # get the unique hotels for this group
    hotels = []
    hotel_nights = []
    transel_agent = Agent.objects.get(name="TRANSEL LTD")
    for h in group_transfer.group_travelday.values('hotel'):
        hotel_id = h['hotel']
        if hotel_id:
            hotel = Hotel.objects.get(id=hotel_id)
            if hotel not in hotels:
                hotels.append(hotel)

    for hotel in hotels:
        hotel_days = [day for day in group_transfer.group_travelday.all() if day.hotel == hotel]
        temp = "%d Nights: " % len(hotel_days)
        temp += "%s" % datetime.date.strftime(hotel_days[0].date, '%d %b %Y')

        if len(hotel_days) > 1:
            if days_are_consecutive([td.date for td in hotel_days]):
                temp += " to %s" % datetime.date.strftime(hotel_days[-1].date + datetime.timedelta(days=1), '%d %b %Y')
            else:
                dates = [datetime.date.strftime(day.date, '%d %b %Y') for day in hotel_days]
                temp = "Nights:\n" + "\n".join(f"- {date} // 1 night" for date in dates)

        hotel_nights.append(temp)

    for i, hotel in enumerate(hotels):
        try:
            rooming_list = DocRoomingList.objects.get(
                group_transfer=group_transfer,
                hotel=hotel
            )
            rooming_list.doc_Date = TravelDay.objects.filter(
                hotel=hotel,
                group_transfer=group_transfer
            ).order_by('date')[0].date
            rooming_list.doc_nights = hotel_nights[i]
            rooming_list.save()
        except DocRoomingList.DoesNotExist:
            rooming_list = DocRoomingList()
            rooming_list.group_transfer = group_transfer
            rooming_list.hotel = hotel
            rooming_list.doc_From = transel_agent.contact
            rooming_list.doc_Date = TravelDay.objects.filter(hotel=hotel, group_transfer=group_transfer).order_by('date')[0].date
            rooming_list.doc_Attn = hotel.contact
            rooming_list.doc_Ref = hotel.contact
            rooming_list.doc_To = hotel.contact

            if group_transfer.roomtext:
                rooming_list.text = group_transfer.roomtext
            else:
                rooming_list.text = ''

            rooming_list.doc_nights = hotel_nights[i]
            rooming_list.save()
    for rl in DocRoomingList.objects.filter(group_transfer=group_transfer):
        if rl.hotel not in hotels:
            rl.delete()


def gen_rooming_list_docx(group, rooming_list, one):
    refcode = group.refcode
    doc = rooming_list
    to = doc.doc_To
    date = doc.doc_Date
    docFrom = doc.doc_From
    attn = doc.doc_Attn
    text = doc.roomtext
    room_desc = doc.room_desc
    meal_desc = doc.meal_desc

    # get details info from Group Transfer table
    group_transfer = GroupTransfer.objects.get(refcode=refcode)

    from django.db import connection
    cursor = connection.cursor()
    cursor.execute("""
        SELECT room_desc,
        meal_desc
        FROM webapp_grouptransfer
        WHERE refcode = '""" + refcode + "';")
    res = cursor.fetchall()
    group_transfer.room_desc = res[0][0]
    group_transfer.meal_desc = res[0][1]

    hotel_name = doc.hotel.name
    current_hotel = doc.hotel
    nights = doc.doc_nights

    if not DocRoomingList.objects.get(
        group_transfer=group,
        hotel=current_hotel
    ):
        doc = DocRoomingList()
        doc.doc_To = to
        doc.doc_Date = date
        doc.doc_From = docFrom
        doc.doc_Attn = attn
        doc.roomtext = text
        doc.group_transfer = group
        doc.hotel = current_hotel
        doc.doc_nights = nights
        doc.save()
        update_roominglists(group)
    else:
        doc = DocRoomingList.objects.get(
            group_transfer=group,
            hotel=current_hotel
        )
        doc.doc_To = to
        doc.doc_Date = date
        doc.doc_From = docFrom
        doc.doc_Attn = attn
        doc.roomtext = text
        doc.doc_nights = nights
        doc.group_transfer = group
        doc.hotel = current_hotel
        doc.save()
        update_roominglists(group)

    logo = str(BASE_DIR) + '/static/images/logos/doc_header.png'
    document = Document()

    # Get the default section of the document
    section = document.sections[0]

    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Set the section start to continuous
    section.start_type = 1

    # Get the header of the section
    header = section.header

    # Create a table in the header with a single row and a single column
    header_table = header.add_table(rows=1, cols=1, width=1900)

    # Set the width of the cell in the header table
    header_table.columns[0].width = Inches(1.5)

    # Add your logo to the cell
    logo_cell = header_table.cell(0, 0)
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    logo_paragraph = logo_cell.paragraphs[0]
    logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture(logo, width=Inches(7.5))

    # Set the header table alignment to center
    header_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add a paragraph after the header
    document.add_paragraph().add_run(hotel_name).bold = True

    # Get the last paragraph in the document
    last_paragraph = document.paragraphs[-1]

    # Set paragraph alignment to center
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set font properties
    font = last_paragraph.runs[0].font
    font.size = Pt(18)
    font.color.rgb = RGBColor(128, 128, 128)  # Grey color

    document.add_paragraph().add_run(date.strftime("%d %b %Y")).bold = True

    # Get the last paragraph in the document
    last_paragraph = document.paragraphs[-1]

    # Set paragraph alignment to center
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set font properties
    font = last_paragraph.runs[0].font
    font.size = Pt(18)
    font.color.rgb = RGBColor(128, 128, 128)  # Grey color

    table = document.add_table(rows=4, cols=2)
    table.style = 'TableGrid'

    for cell in table.columns[0].cells:
        cell.width = Inches(1.5)

    hdr_cells = table.rows[0].cells

    hdr_cells[0].text = 'Group Refcode'
    hdr_cells[1].text = refcode

    # hdr_cells = table.rows[1].cells
    # hdr_cells[0].text = 'Hotel Name'
    # hdr_cells[1].text = hotel_name

    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = 'Rooms Description'

    if room_desc == '' or meal_desc is None:
        hdr_cells[1].text = str(group.room_desc)
    else:
        hdr_cells[1].text = str(room_desc)

    hdr_cells = table.rows[2].cells
    hdr_cells[0].text = 'Nights'
    hdr_cells[1].text = nights

    hdr_cells = table.rows[3].cells
    hdr_cells[0].text = 'Meal Description'

    if meal_desc == '' or meal_desc is None:
        hdr_cells[1].text = str(group.meal_desc)
    else:
        hdr_cells[1].text = str(meal_desc)

    document.add_paragraph("\n")

    soup = BeautifulSoup(text)
    # TBD : need to double check is the table or not, and how many tables

    for row in table.rows:
        for cell in row.cells:

            # Set vertical alignment of cell content to center
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # Access the paragraph in the cell and apply formatting
            paragraph = cell.paragraphs[0]

            # Set font color to red
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(75, 75, 75)

            # Set font size to 14
            paragraph.style.font.size = Pt(14)

    # Add a paragraph after the header
    document.add_paragraph().add_run("Rooming List").bold = True

    # Get the last paragraph in the document
    last_paragraph = document.paragraphs[-1]

    # Set paragraph alignment to center
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set font properties
    font = last_paragraph.runs[0].font
    font.size = Pt(18)
    font.color.rgb = RGBColor(128, 128, 128)  # Grey color

    try:
        table = soup.findAll('table')
        tab = table[0]
        rows = tab.findAll('tr')
        try:
            textTable = document.add_table(
                rows=len(rows),
                cols=len(rows[0].findAll('td'))
            )
            textTable.style = 'Table Grid'

            for cell in textTable.columns[0].cells:
                cell.width = Inches(0.8)
            for cell in textTable.columns[1].cells:
                cell.width = Inches(0.3)
            for cell in textTable.columns[2].cells:
                cell.width = Inches(3)
            for cell in textTable.columns[3].cells:
                cell.width = Inches(0.2)
            for cell in textTable.columns[4].cells:
                cell.width = Inches(1.0)

        except Exception as exc:
            print(exc)
        i = 0

        # Fill all the cells and dont care about row spam
        for tr in tab.findAll('tr'):
            hdr_cells = textTable.rows[i].cells
            j = 0
            for td in tr.findAll('td'):
                # When met "rowspan", move right and add merge cells
                if td.has_key("rowspan"):
                    for k in range(1, int(td['rowspan'])):
                        hdr_cells[j].merge(textTable.cell(i + k, j))
                if hdr_cells[j].text:
                    j = j + 1
                cell_text = td.getText().replace('&nbsp;', '').strip()  # Remove leading and trailing whitespace
                hdr_cells[j].text = str(cell_text)
                hdr_cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                paragraph = hdr_cells[j].paragraphs[0]
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(60, 60, 60)
                paragraph.style.font.size = Pt(10)

                j = j + 1
            i = i + 1

        # if note is not None:
        #     document.add_paragraph().add_run(f'Notes: \n {str(note)}').bold = True

        #     # Get the last paragraph in the document
        #     last_paragraph = document.paragraphs[-1]

        #     # Set paragraph alignment to center
        #     last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        #     # Set font properties
        #     font = last_paragraph.runs[0].font
        #     font.size = Pt(18)
        #     font.color.rgb = RGBColor(128, 128, 128)  # Grey color

    except Exception as exc:
        print(exc)

    # Create a footer
    footer = document.sections[0].footer

    # Add a paragraph to the footer with company details
    footer_paragraph = footer.paragraphs[0]
    if group.refcode.startswith("COL"):
        footer_paragraph.text = "Cosmoplan International Travel | 105-109 Sumatra Road, NW6 1PL | Phone: +44 20 81436880 | Email: info@cosmoplan.co.uk"
    else:
        footer_paragraph.text = "Cosmoplan International Travel | Xenofontos 10, Athens, Greece | Phone: +30 2109219400 | Email: info@cosmoplan.gr"
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    fileName = str(BASE_DIR) + '/files/rooming_lists/' + refcode + '/RoomingList_' + refcode + '_' + str(current_hotel.id) + '.docx'

    folder_path = os.path.join(BASE_DIR, 'files', 'rooming_lists', refcode)
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)

    document.save(fileName)
    return fileName


class ChangeRoomingListRoomText(generics.UpdateAPIView):
    """
    URL: change_rooming_list_room_text/(?P<refcode>.*)$
    Descr: Changes rooming list's roomtext. Room text is an HTML field
    """

    # Cross-Site Request Forgery
    @csrf_exempt
    def post(self, request, refcode):
        context = {"errormsg": ''}
        token_str = request.headers['User-Token']
        user = get_user(token_str)

        # Permission
        if not can_update(user.id, 'GT'):
            context = {"errormsg": "You do not have permission to update a Group's rooming list."}
            return Response(status=401, data=context)

        # Get rooming list
        rooming_list = DocRoomingList.objects.get(id=request.data['rooming_list_id'])

        # Get previous room text for logging
        prev_room_text = rooming_list.roomtext

        # Get new roomtext
        room_text = request.data['room_text'].strip()
        rooming_list.roomtext = room_text
        rooming_list.save()
        History.objects.create(
            user=user,
            model_name='GT',
            action='CRE',
            description=f"User updated rooming list's text of group : {refcode} with date of : \
                {rooming_list.doc_Date} from {prev_room_text} to {room_text}"
        )
        context['rooming_list'] = RoomingListSerializer(rooming_list).data
        return Response(data=context, status=200)


class ChangeRoomingListRoomDescr(generics.ListCreateAPIView):
    """
    URL: change_rooming_list_room_description/(?P<refcode>.*)$
    Descr: Changes rooming lists room description. Free text field
    """

    # Cross-Site Request Forgery
    @csrf_exempt
    def post(self, request, refcode):
        context = {"errormsg": ''}
        token_str = request.headers['User-Token']
        user = get_user(token_str)

        # Permission
        if not can_update(user.id, 'GT'):
            context = {"errormsg": "You do not have permission to update a Group's rooming list room description."}
            return Response(status=401, data=context)

        # Get rooming list
        rooming_list = DocRoomingList.objects.get(id=request.data['rooming_list_id'])

        # Get previous room desc for logging
        prev_room_desc = rooming_list.room_desc if rooming_list.room_desc is not None else 'N/A'

        # Get new room desc
        new_room_desc = request.data['room_desc'].strip()
        try:
            rooming_list.room_desc = new_room_desc
            rooming_list.save()
            History.objects.create(
                user=user,
                model_name='GT',
                action='UPD',
                description=f"User updated rooming list's room description of group : {refcode} with date of : \
                    {rooming_list.doc_Date} from {prev_room_desc} to {new_room_desc}"
            )
            context['rooming_list'] = RoomingListSerializer(rooming_list).data
        except Exception as a:
            context['errormsg'] = a
            return Response(data=context, status=400)
        return Response(data=context, status=200)


class ChangeRoomingListMealDescr(generics.ListCreateAPIView):
    """
    URL: change_rooming_list_meal_description/(?P<refcode>.*)$
    Descr: Changes rooming lists meal description. Free text field
    """

    # Cross-Site Request Forgery
    @csrf_exempt
    def post(self, request, refcode):
        context = {"errormsg": ''}
        token_str = request.headers['User-Token']
        user = get_user(token_str)

        # Permission
        if not can_update(user.id, 'GT'):
            context = {"errormsg": "You do not have permission to update a Group's rooming list meal description."}
            return Response(status=401, data=context)

        # Get rooming list
        rooming_list = DocRoomingList.objects.get(id=request.data['rooming_list_id'])

        # Get previous meal desc for logging
        prev_meal_desc = rooming_list.meal_desc if rooming_list.meal_desc is not None else 'N/A'

        # Get new meal desc
        new_meal_desc = request.data['meal_desc'].strip()

        try:
            rooming_list.meal_desc = new_meal_desc
            rooming_list.save()
            History.objects.create(
                user=user,
                model_name='GT',
                action='UPD',
                description=f"User updated rooming list's meal description of group : {refcode} with date of : \
                    {rooming_list.doc_Date} from {prev_meal_desc} to {new_meal_desc}"
            )
            context['rooming_list'] = RoomingListSerializer(rooming_list).data
        except Exception as a:
            context['errormsg'] = a
            return Response(data=context, status=400)
        return Response(data=context, status=200)


class ChangeRoomingListNote(generics.ListCreateAPIView):
    """
    URL: change_rooming_list_note/(?P<refcode>.*)$
    Descr: Changes rooming lists note. Free text field
    """

    # Cross-Site Request Forgery
    @csrf_exempt
    def post(self, request, refcode):
        context = {"errormsg": ''}
        token_str = request.headers['User-Token']
        user = get_user(token_str)

        # Permission
        if not can_update(user.id, 'GT'):
            context = {"errormsg": "You do not have permission to update a Group's rooming list note."}
            return Response(status=401, data=context)

        # Get rooming list
        rooming_list = DocRoomingList.objects.get(id=request.data['rooming_list_id'])

        # Get previous note for logging
        prev_note = rooming_list.note if rooming_list.note is not None else 'N/A'

        # Get new note
        new_note = request.data['note'].strip()
        try:
            rooming_list.note = new_note
            rooming_list.save()
            History.objects.create(
                user=user,
                model_name='GT',
                action='UPD',
                description=f"User updated rooming list's note of group : {refcode} with date of : \
                    {rooming_list.doc_Date} from {prev_note} to {new_note}"
            )
            context['rooming_list'] = RoomingListSerializer(rooming_list).data
        except Exception as a:
            context['errormsg'] = a
            return Response(data=context, status=400)
        return Response(data=context, status=200)


class EditMassRoomingListFields(generics.ListCreateAPIView):
    """
    URL: edit_mass_rooming_list_fields/(?P<refcode>.*)$
    Descr: Edits note, meal and room descriptions
    for all rooming lists of a group
    """

    # Cross-Site Request Forgery
    @csrf_exempt
    def post(self, request, refcode):
        context = {"errormsg": ''}
        token_str = request.headers['User-Token']
        user = get_user(token_str)

        # Permission
        if not can_update(user.id, 'GT'):
            context = {"errormsg": "You do not have permission to update a Group's rooming list."}
            return Response(status=401, data=context)

        # Get group
        group = GroupTransfer.objects.get(refcode=refcode)

        # Get selected hotels
        selected_hotels = request.data['selected_hotels']
        hotel_ids = [Hotel.objects.get(name=i).id for i in selected_hotels if selected_hotels[i]]

        # Get their rooming lists
        rooming_lists = DocRoomingList.objects.filter(hotel_id__in=hotel_ids, group_transfer_id=group.id)

        for rooming_list in rooming_lists:
            rooming_list.note = request.data['all_note']
            rooming_list.roomtext = request.data['room_text']
            rooming_list.save()
            History.objects.create(
                user=user,
                model_name='GT',
                action='UPD',
                description=f"User updated rooming lists of group: {group.refcode} \
                fields: room desc, meal desc, notes."
            )
        group.roomtext = request.data['room_text']
        group.save()
        context['group'] = GroupSerializer(group).data

        History.objects.create(
            user=user,
            model_name='GT',
            action='UPD',
            description=f"User : {user.username} updated group's Rooming List text "
        )

        return Response(data=context, status=200)


class SendAllRoomingLists(generics.UpdateAPIView):
    """
    URL: send_all_rooming_lists/(?P<refcode>.*)$
    Descr: Sends all rooming lists to hotels
    """

    # Cross-Site Request Forgery
    @csrf_exempt
    def post(self, request, refcode):
        context = {"errormsg": '', "new_refcode": ''}
        token_str = request.headers['User-Token']
        user = get_user(token_str)

        # Permission
        if not can_update(user.id, 'GT'):
            context = {"errormsg": "You do not have permission to send a Rooming List"}
            return Response(status=401, data=context)

        # Send from will always have the user's email
        send_from = request.data['send_from']
        user_profile = UserProfile.objects.get(user_id=user.id)

        # Get group
        group = GroupTransfer.objects.get(refcode=refcode)

        # Get group's rooming lists
        rooming_lists = DocRoomingList.objects.filter(group_transfer_id=group.id)

        recipients = []

        # loop over rooming lists
        for rooming_list in rooming_lists:

            # Just for test purposes.
            gen_rooming_list_docx(rooming_list.group_transfer, rooming_list, False)

            subject = f'Cosmoplan Rooming List for {group.refcode} {rooming_list.doc_Date}'
            message = f"""
                <p>Dear, </p>
                <p>
                    I am writing to inform you that the rooming list for the upcoming group reservation is now finalized.
                    <br/>
                    Kindly find the details below:
                </p>
                <table class=MsoNormalTable border=1 align=left style='margin-left:6.75pt;'>
                    <tr>
                        <td>
                            <p>
                                <b>
                                    <span>
                                        GROUP REF :
                                    </span>
                                </b>
                            </p>
                        </td>
                        <td>
                            <p>
                                <b>
                                    <span>
                                        {group.refcode}
                                    </span>
                                </b>
                            </p>
                        </td>
                    </tr>
                    <tr>
                        <td>
                            <p>
                                <b>
                                    <span>
                                        Hotel Name :
                                    </span>
                                </b>
                            </p>
                        </td>
                        <td>
                            <p>
                                <b>
                                    <span>
                                        {rooming_list.hotel.name}
                                    </span>
                                </b>
                            </p>
                        </td>
                    </tr>
                    <tr>
                        <td>
                            <p>
                                <b>
                                    <span>
                                        Check in date :
                                    </span>
                                </b>
                            </p>
                        </td>
                        <td>
                            <p>
                                <b>
                                    <span>
                                        {rooming_list.doc_Date}
                                    </span>
                                </b>
                            </p>
                        </td>
                    </tr>

                    <tr>
                        <td>
                            <p>
                                <b>
                                    <span>
                                        Rooms :
                                    </span>
                                </b>
                            </p>
                        </td>
                        <td>
                            <p>
                                <b>
                                    <span>
                                        {group.room_desc}
                                    </span>
                                </b>
                            </p>
                        </td>
                    </tr>

                    <tr>
                        <td>
                            <p>
                                <b>
                                    <span>
                                        Night(s) :
                                    </span>
                                </b>
                            </p>
                        </td>
                        <td>
                            <p>
                                <b>
                                    <span>
                                        {rooming_list.doc_nights}
                                    </span>
                                </b>
                            </p>
                        </td>
                    </tr>

                    <tr>
                        <td>
                            <p>
                                <b>
                                    <span>
                                        Meal Description :
                                    </span>
                                </b>
                            </p>
                        </td>
                        <td>
                            <p>
                                <b>
                                    <span>
                                        {group.meal_desc}
                                    </span>
                                </b>
                            </p>
                        </td>
                    </tr>
            """
            if rooming_list.note is not None and rooming_list.note != "" and rooming_list.note != 'N/A':
                message += f"""
                    <tr>
                        <td>
                            <p>
                                <b>
                                    <span>
                                        Note :
                                    </span>
                                </b>
                            </p>
                        </td>
                        <td>
                            <p>
                                <b>
                                    <span>
                                        {str(rooming_list.note)}
                                    </span>
                                </b>
                            </p>
                        </td>
                    </tr>
                </table>
                <br/>
                <br/>
                <br/>
                <br/>
                <br/>
                <br/>
                <br/>
                <br/>
                <br/>
                <br/>
                <br/>
            """

            else:
                message += """
                </table>
                <br/>
                <br/>
                <br/>
                <br/>
                <br/>
                <br/>
                <br/>
                <br/>
                <br/>
                <br/>
            """

            message += """
                <p> Thank you for your prompt attention to this matter. <p/>
                <p>Should you require any further information or clarification, please do not hesitate to contact us. <p/>
                <p> Looking forward to your swift response. <p/>
                <br/>
            """

            if send_from == user.email:
                message += user_profile.signature
            else:
                if user_profile.secondary_signature is not None:
                    message += user_profile.secondary_signature

            # User for sweet alert 2
            recipients.append(Contact.objects.get(id=rooming_list.doc_To_id).email)
            send_to = [Contact.objects.get(id=rooming_list.doc_To_id).email]
            msg = EmailMultiAlternatives(subject, message, send_from, send_to, cc=[send_from])
            msg.attach_alternative(message, "text/html")
            msg.content_subtype = 'html'

            # Attach the file
            attachment = open(str(BASE_DIR) + '/files/rooming_lists/' + refcode + '/RoomingList_' + refcode + '_' + str(rooming_list.hotel.id) + '.docx', 'rb')
            msg.attach('RoomingList_' + refcode + '_' + str(rooming_list.hotel.id) + '.docx', attachment.read(), 'text/plain')

            # Set the content type of the email to HTML
            msg.attach_alternative(message, "text/html")

            if not rooming_list.sent:
                msg.send()
                rooming_list.sent = True
                rooming_list.save()
            context = {"recipients": recipients}

        History.objects.create(
            user=user,
            model_name='GT',
            action='UPD',
            description=f"User : {user.username} sent all rooming lists of group: {group.refcode}. Recipients: {recipients}"
        )

        return Response(data=context, status=200)


class SendRoomingList(generics.UpdateAPIView):
    """
    URL: send_all_rooming_lists/(?P<refcode>.*)$
    Descr: Sends rooming list to specific hotel
    """

    # Cross-Site Request Forgery
    @csrf_exempt
    def post(self, request, refcode):
        context = {"errormsg": '', "new_refcode": ''}
        token_str = request.headers['User-Token']
        user = get_user(token_str)

        # Get group
        group = GroupTransfer.objects.get(refcode=refcode)

        # Permission
        if not can_update(user.id, 'GT'):
            context = {"errormsg": "You do not have permission to send a Rooming List"}
            return Response(status=401, data=context)

        # Get rooming list id
        rooming_list_id = request.data['rooming_list_id']

        try:
            rooming_list = DocRoomingList.objects.get(id=rooming_list_id)
        except DocRoomingList.DoesNotExist:
            pass

        gen_rooming_list_docx(rooming_list.group_transfer, rooming_list, True)

        user_profile = UserProfile.objects.get(user=User.objects.get(username=user.username))

        recipients = [Contact.objects.get(id=rooming_list.doc_To_id).email, ]

        # Send from will always have the user's email
        send_from = request.data['send_from']
        subject = f'Cosmoplan Rooming List for {group.refcode} {rooming_list.doc_Date}'
        message = f"""
            <p>Dear, </p>
            <p>
                I am writing to inform you that the rooming list for the upcoming group reservation is now finalized.
                <br/>
                Kindly find the details below:
            </p>
            <table class=MsoNormalTable border=1 align=left style='margin-left:6.75pt;'>
                <tr>
                    <td>
                        <p>
                            <b>
                                <span>
                                    GROUP REF :
                                </span>
                            </b>
                        </p>
                    </td>
                    <td>
                        <p>
                            <b>
                                <span>
                                    {group.refcode}
                                </span>
                            </b>
                        </p>
                    </td>
                </tr>
                <tr>
                    <td>
                        <p>
                            <b>
                                <span>
                                    Hotel Name :
                                </span>
                            </b>
                        </p>
                    </td>
                    <td>
                        <p>
                            <b>
                                <span>
                                    {rooming_list.hotel.name}
                                </span>
                            </b>
                        </p>
                    </td>
                </tr>
                <tr>
                    <td>
                        <p>
                            <b>
                                <span>
                                    Check in date :
                                </span>
                            </b>
                        </p>
                    </td>
                    <td>
                        <p>
                            <b>
                                <span>
                                    {rooming_list.doc_Date}
                                </span>
                            </b>
                        </p>
                    </td>
                </tr>

                <tr>
                    <td>
                        <p>
                            <b>
                                <span>
                                    Rooms :
                                </span>
                            </b>
                        </p>
                    </td>
                    <td>
                        <p>
                            <b>
                                <span>
                                    {group.room_desc}
                                </span>
                            </b>
                        </p>
                    </td>
                </tr>

                <tr>
                    <td>
                        <p>
                            <b>
                                <span>
                                    Night(s) :
                                </span>
                            </b>
                        </p>
                    </td>
                    <td>
                        <p>
                            <b>
                                <span>
                                    {rooming_list.doc_nights}
                                </span>
                            </b>
                        </p>
                    </td>
                </tr>

                <tr>
                    <td>
                        <p>
                            <b>
                                <span>
                                    Meal Description :
                                </span>
                            </b>
                        </p>
                    </td>
                    <td>
                        <p>
                            <b>
                                <span>
                                    {group.meal_desc}
                                </span>
                            </b>
                        </p>
                    </td>
                </tr>
            </table>
            <br/>
            <br/>
            <br/>
            <br/>
            <br/>
            <br/>
            <br/>
            <br/>
            <br/>
            <br/>
            <p> Thank you for your prompt attention to this matter. <p/>
            <p>Should you require any further information or clarification, please do not hesitate to contact us. <p/>
            <p> Looking forward to your swift response. <p/>
            <br/>
        """

        if send_from == user.email:
            message += user_profile.signature
        else:
            if user_profile.secondary_signature is not None:
                message += user_profile.secondary_signature

        send_to = [Contact.objects.get(id=rooming_list.doc_To_id).email]
        msg = EmailMultiAlternatives(subject, message, send_from, send_to, cc=[send_from])
        msg.attach_alternative(message, "text/html")
        msg.content_subtype = 'html'

        # Attach the file
        attachment = open(str(BASE_DIR) + '/files/rooming_lists/' + refcode + '/RoomingList_' + refcode + '_' + str(rooming_list.hotel.id) + '.docx', 'rb')
        msg.attach('RoomingList_' + refcode + '_' + str(rooming_list.hotel.id) + '.docx', attachment.read(), 'text/plain')

        # Set the content type of the email to HTML
        msg.attach_alternative(message, "text/html")

        History.objects.create(
            user=user,
            model_name='GT',
            action='UPD',
            description=f"User : {user.username} sent rooming list of hotel: {rooming_list.hotel.contact.email} and group: \
                {group.refcode}."
        )

        rooming_list.sent = True
        rooming_list.save()

        msg.send()
        context = {"recipients": recipients}

        return Response(data=context, status=200)


class ToggleCheckmark(generics.ListCreateAPIView):
    """
    URL: toggle_checkmark/
    Descr: Changes sent status of a rooming list.
    """

    # Cross-Site Request Forgery
    @csrf_exempt
    def post(self, request):
        context = {"errormsg": ''}
        token_str = request.headers['User-Token']
        user = get_user(token_str)

        # Permission
        if not can_update(user.id, 'GT'):
            context = {"errormsg": "You do not have permission to update a Group's rooming list."}
            return Response(status=401, data=context)

        group = GroupTransfer.objects.get(refcode=request.data['refcode'])

        rooming_list = DocRoomingList.objects.get(id=request.data['rl_id'])
        rooming_list.sent = not request.data['checked']

        rooming_list.save()
        History.objects.create(
            user=user,
            model_name='GT',
            action='UPD',
            description=f"User updated rooming list's sent status to {request.data['checked']}"
        )
        context['group'] = GroupSerializer(group).data
        return Response(data=context, status=200)
